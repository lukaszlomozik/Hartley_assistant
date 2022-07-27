import os
import docx


class Report:

    def create_file(self):
        self.doc = docx.Document()

    def add_heading(self, text):
        self.doc.add_heading(text, 0)

    def add_text(self, text, bold=False):
        p = self.doc.add_paragraph()
        p.add_run(text).bold = bold

    def add_picture(self, path):
        self.doc.add_picture(path)
        os.remove(path)

    def add_table(self, columns, values, rd):
        rows = len(values[0])
        cols = len(columns)
        table = self.doc.add_table(rows=rows + 1, cols=cols)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i in range(cols):
            hdr_cells[i].text = columns[i]
        for n in range(cols):
            for i in range(rows):
                hdr_cells = table.rows[i + 1].cells
                hdr_cells[n].text = str(round(values[n][i], rd))

    def save_file(self, file):
        if os.path.isfile(file):
            os.remove(file)
        self.doc.save(file)
